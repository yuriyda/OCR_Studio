"""Тесты конвертера markdown → docx (через html_to_docx walker)."""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.document import Document as DocxDocument

from app.converters import html_to_docx, md_to_docx


def _open_docx(data: bytes) -> DocxDocument:
    return Document(io.BytesIO(data))


def test_html_to_docx_simple_paragraph():
    data = html_to_docx("<p>hello world</p>")
    doc = _open_docx(data)
    paras = [p.text for p in doc.paragraphs if p.text]
    assert paras == ["hello world"]


def test_html_to_docx_h1_h2_h3():
    data = html_to_docx("<h1>Title</h1><h2>Section</h2><h3>Sub</h3>")
    doc = _open_docx(data)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 3
    assert headings[0].text == "Title"
    assert headings[0].style.name == "Heading 1"
    assert headings[1].style.name == "Heading 2"
    assert headings[2].style.name == "Heading 3"


def test_html_to_docx_h4_h5_h6_clamped_to_4():
    data = html_to_docx("<h6>Deep</h6>")
    doc = _open_docx(data)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 1
    assert headings[0].style.name in ("Heading 4",)


def test_html_to_docx_multiple_paragraphs_keep_order():
    data = html_to_docx("<p>first</p><p>second</p><p>third</p>")
    doc = _open_docx(data)
    paras = [p.text for p in doc.paragraphs if p.text]
    assert paras == ["first", "second", "third"]


def test_html_to_docx_empty_input_returns_valid_docx():
    data = html_to_docx("")
    doc = _open_docx(data)
    assert isinstance(doc, DocxDocument)


def test_html_to_docx_unordered_list():
    data = html_to_docx("<ul><li>alpha</li><li>beta</li><li>gamma</li></ul>")
    doc = _open_docx(data)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullets) == 3
    assert [p.text for p in bullets] == ["alpha", "beta", "gamma"]


def test_html_to_docx_ordered_list():
    data = html_to_docx("<ol><li>one</li><li>two</li></ol>")
    doc = _open_docx(data)
    nums = [p for p in doc.paragraphs if p.style.name == "List Number"]
    assert len(nums) == 2
    assert [p.text for p in nums] == ["one", "two"]


def test_html_to_docx_list_then_paragraph():
    data = html_to_docx("<ul><li>x</li><li>y</li></ul><p>after</p>")
    doc = _open_docx(data)
    visible = [(p.style.name, p.text) for p in doc.paragraphs if p.text]
    assert visible == [
        ("List Bullet", "x"),
        ("List Bullet", "y"),
        ("Normal", "after"),
    ]


def test_html_to_docx_inline_bold():
    data = html_to_docx("<p>plain <strong>bold</strong> text</p>")
    doc = _open_docx(data)
    p = [p for p in doc.paragraphs if p.text][0]
    runs = list(p.runs)
    bold_runs = [r for r in runs if r.bold]
    assert bold_runs, "no bold run found"
    assert "bold" in "".join(r.text for r in bold_runs)


def test_html_to_docx_inline_italic():
    data = html_to_docx("<p>see <em>this</em> word</p>")
    doc = _open_docx(data)
    p = [p for p in doc.paragraphs if p.text][0]
    italic_runs = [r for r in p.runs if r.italic]
    assert italic_runs and "this" in "".join(r.text for r in italic_runs)


def test_html_to_docx_inline_code():
    data = html_to_docx("<p>use <code>print()</code> here</p>")
    doc = _open_docx(data)
    p = [p for p in doc.paragraphs if p.text][0]
    code_runs = [r for r in p.runs if r.font.name == "Courier New"]
    assert code_runs and "print()" in "".join(r.text for r in code_runs)


def test_html_to_docx_link_renders_text_and_url():
    data = html_to_docx('<p>see <a href="https://example.com">site</a> docs</p>')
    doc = _open_docx(data)
    p = [p for p in doc.paragraphs if p.text][0]
    text = p.text
    assert "site" in text
    assert "https://example.com" in text  # URL appears (in parens or as raw)


def test_html_to_docx_code_block():
    html = "<pre><code>line1\nline2\nline3</code></pre>"
    data = html_to_docx(html)
    doc = _open_docx(data)
    code_paras = [p for p in doc.paragraphs if any(r.font.name == "Courier New" for r in p.runs)]
    assert code_paras, "no code paragraph"
    assert "line1" in code_paras[0].text
    assert "line3" in code_paras[0].text


def test_html_to_docx_blockquote():
    data = html_to_docx("<blockquote><p>quoted text</p></blockquote>")
    doc = _open_docx(data)
    quotes = [p for p in doc.paragraphs if p.style.name in ("Intense Quote", "Quote")]
    assert quotes and "quoted text" in quotes[0].text


def test_html_to_docx_hr_creates_separator_paragraph():
    data = html_to_docx("<p>before</p><hr/><p>after</p>")
    doc = _open_docx(data)
    paras = [p.text for p in doc.paragraphs]
    assert "before" in paras
    assert "after" in paras
    assert paras.index("before") < paras.index("after")


def test_html_to_docx_table_basic_2x2():
    html = """
    <table>
      <thead><tr><th>A</th><th>B</th></tr></thead>
      <tbody><tr><td>1</td><td>2</td></tr></tbody>
    </table>
    """
    data = html_to_docx(html)
    doc = _open_docx(data)
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 2 and len(t.columns) == 2
    assert t.rows[0].cells[0].text == "A"
    assert t.rows[0].cells[1].text == "B"
    assert t.rows[1].cells[0].text == "1"
    assert t.rows[1].cells[1].text == "2"


def test_html_to_docx_table_uneven_rows_padded():
    html = "<table><tr><td>a</td><td>b</td><td>c</td></tr><tr><td>x</td></tr></table>"
    data = html_to_docx(html)
    doc = _open_docx(data)
    t = doc.tables[0]
    assert len(t.columns) == 3  # max width
    assert t.rows[1].cells[0].text == "x"
