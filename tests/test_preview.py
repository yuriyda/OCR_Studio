"""Tests for app/preview.py."""
from app import preview


def test_markdown_headings():
    html = preview.markdown_to_html("# Title\n## Sub")
    assert "<h1>Title</h1>" in html
    assert "<h2>Sub</h2>" in html


def test_markdown_table():
    md = "| A | B |\n| --- | --- |\n| 1 | 2 |"
    html = preview.markdown_to_html(md)
    assert "<table>" in html
    assert "<th>A</th>" in html
    assert "<td>1</td>" in html


def test_markdown_lists():
    html = preview.markdown_to_html("- a\n- b")
    assert "<ul>" in html
    assert "<li>a</li>" in html


def test_markdown_fenced_code():
    html = preview.markdown_to_html("```\ncode\n```")
    assert "<code>" in html
    assert "code" in html


def test_markdown_emphasis():
    html = preview.markdown_to_html("*em* and **bold**")
    assert "<em>em</em>" in html
    assert "<strong>bold</strong>" in html


def test_sanitize_strips_script():
    bad = "<p>hi</p><script>alert(1)</script>"
    clean = preview.sanitize_html(bad)
    assert "<script" not in clean
    assert "<p>hi</p>" in clean


def test_sanitize_strips_iframe():
    clean = preview.sanitize_html('<iframe src="x"></iframe><p>ok</p>')
    assert "<iframe" not in clean


def test_sanitize_strips_event_handlers():
    clean = preview.sanitize_html('<p onerror="alert(1)">x</p>')
    assert "onerror" not in clean


def test_sanitize_strips_javascript_href():
    clean = preview.sanitize_html('<a href="javascript:alert(1)">x</a>')
    assert "javascript:" not in clean


def test_sanitize_keeps_table_attrs():
    clean = preview.sanitize_html('<table><tr><td colspan="2">x</td></tr></table>')
    assert 'colspan="2"' in clean


def test_text_to_html_wraps_in_pre():
    clean = preview.text_to_html("a\nb\n<script>")
    assert "<pre>" in clean
    assert "&lt;script&gt;" in clean
    assert "<script>" not in clean


def test_docx_to_html_table_preserved(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("Hi", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    docx_path = tmp_path / "x.docx"
    doc.save(docx_path)

    html = preview.docx_to_html(docx_path.read_bytes())
    assert "<table>" in html
    assert "A" in html and "1" in html
    assert "<h1>" in html or "Hi" in html


def test_docx_to_html_sanitized():
    from docx import Document
    import io as _io
    d = Document()
    d.add_paragraph("text")
    buf = _io.BytesIO()
    d.save(buf)
    html = preview.docx_to_html(buf.getvalue())
    assert "<script" not in html


def test_block_formula_passthrough():
    """markdown_to_html must preserve $$...$$ math blocks unchanged."""
    from app.preview import markdown_to_html
    md = "Before paragraph\n\n$$\n\\frac{x^2}{y}\n$$\n\nAfter paragraph"
    html = markdown_to_html(md)
    assert "\\frac" in html
    assert "$$" in html or "frac{x^2}{y}" in html


def test_multiple_formulas_in_one_doc():
    from app.preview import markdown_to_html
    md = (
        "$$\nE = mc^2\n$$\n\n"
        "Paragraph in between.\n\n"
        "$$\n\\int_0^1 x\\,dx\n$$\n"
    )
    html = markdown_to_html(md)
    assert "mc^2" in html
    assert "\\int" in html


def test_invalid_latex_kept_as_text():
    """Garbled formula content must not crash the renderer; source preserved."""
    from app.preview import markdown_to_html
    md = "$$\n\\frac{ no closing brace\n$$"
    html = markdown_to_html(md)
    assert "no closing brace" in html
