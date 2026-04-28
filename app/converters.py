"""
OCR result converters: markdown → plain text, markdown → docx.

Maintenance notes:
- md_to_txt: only basic regex strips (headings, bold/italic).
- md_to_docx: rendered via markdown→HTML→html_to_docx walker. HTML is parsed by
  BeautifulSoup; DOM walk builds python-docx structures. Inline parsing (bold/italic/
  code/links) is delegated to the markdown library — do not duplicate.
- Do not add dependencies — bs4, markdown, python-docx are already in requirements.
"""
from __future__ import annotations

import io
import re

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt

from . import preview as _preview

_FORMULA_BLOCK_RE = re.compile(r'\$\$\s*\n?(.+?)\n?\s*\$\$', re.DOTALL)


def _flatten_formulas(md: str) -> str:
    """Replace block math `$$...$$` with backtick-wrapped `[formula: ...]` placeholders.

    Preserves LaTeX source verbatim so users can copy it into Equation Editor.
    Backticks make the placeholder render in monospace via the existing code-walker.
    """
    return _FORMULA_BLOCK_RE.sub(
        lambda m: f"`[formula: {m.group(1).strip()}]`",
        md,
    )


def md_to_txt(md: str) -> str:
    """Strip markdown formatting, keep tables as plain text."""
    lines = md.split("\n")
    out = []
    for line in lines:
        line = re.sub(r"^#{1,6}\s+", "", line)
        line = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", line)
        out.append(line)
    return "\n".join(out)


_HEADING_TAGS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 4, "h6": 4}


def html_to_docx(html: str) -> bytes:
    """Walk sanitized HTML → python-docx structures → bytes.

    Supported elements: h1-h6, p, lists, inline formatting, code blocks, blockquotes, hr, tables.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    soup = BeautifulSoup(html or "", "html.parser")
    for child in soup.children:
        _walk_block(doc, child)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _walk_block(doc, node):
    """Walk top-level (block) HTML node → docx paragraph/heading."""
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            doc.add_paragraph(text)
        return
    name = (node.name or "").lower()
    if name in _HEADING_TAGS:
        h = doc.add_heading("", level=_HEADING_TAGS[name])
        _walk_inline(h, node)
        return
    if name == "p":
        para = doc.add_paragraph()
        _walk_inline(para, node)
        return
    if name in ("ul", "ol"):
        style_name = "List Bullet" if name == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            para = doc.add_paragraph(style=style_name)
            _walk_inline(para, li)
        return
    if name == "pre":
        # Markdown fenced ```code``` → <pre><code>...</code></pre>
        text = node.get_text("\n")  # preserve newlines from <code> children
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Courier New"
        return
    if name == "blockquote":
        for sub in node.children:
            if isinstance(sub, NavigableString):
                t = str(sub).strip()
                if t:
                    doc.add_paragraph(t, style="Intense Quote")
                continue
            sub_name = (sub.name or "").lower()
            if sub_name == "p":
                p = doc.add_paragraph(style="Intense Quote")
                _walk_inline(p, sub)
        return
    if name == "hr":
        # Visual separator — empty paragraph; python-docx has no dedicated horizontal rule.
        doc.add_paragraph("")
        return
    if name == "table":
        rows = node.find_all("tr")
        if not rows:
            return
        rows_cells = [r.find_all(["td", "th"]) for r in rows]
        num_cols = max(len(r) for r in rows_cells) if rows_cells else 0
        if num_cols == 0:
            return
        table = doc.add_table(rows=len(rows_cells), cols=num_cols)
        table.style = "Table Grid"
        for ri, cells in enumerate(rows_cells):
            for ci, cell in enumerate(cells):
                if ci < num_cols:
                    table.rows[ri].cells[ci].text = cell.get_text().strip()
        doc.add_paragraph("")  # spacer after table
        return
    # Other block tags handled in later tasks.


def _walk_inline(para, node):
    """Walk inline children of a block node, building python-docx runs."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                para.add_run(text)
            continue
        cname = (child.name or "").lower()
        if cname in ("strong", "b"):
            run = para.add_run(child.get_text())
            run.bold = True
        elif cname in ("em", "i"):
            run = para.add_run(child.get_text())
            run.italic = True
        elif cname == "code":
            run = para.add_run(child.get_text())
            run.font.name = "Courier New"
        elif cname == "a":
            href = child.get("href", "")
            text = child.get_text()
            # MVP: text + URL in parentheses. A real OOXML hyperlink is a separate task.
            if href:
                para.add_run(f"{text} ({href})")
            else:
                para.add_run(text)
        elif cname == "br":
            para.add_run().add_break()
        else:
            # Unknown inline → fallback to text
            para.add_run(child.get_text())


def md_to_docx(md: str) -> bytes:
    """Convert markdown to a .docx file. Returns bytes.

    Pipeline: _flatten_formulas → markdown → HTML (preview.markdown_to_html) → docx (html_to_docx).
    Block math `$$...$$` is replaced with monospace placeholders before HTML rendering.
    """
    md = _flatten_formulas(md or "")
    html = _preview.markdown_to_html(md)
    return html_to_docx(html)
